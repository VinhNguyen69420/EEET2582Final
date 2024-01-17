from io import BytesIO
from docx import Document
from docx.shared import Pt
import xml.etree.ElementTree as ET
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK
from difflib import SequenceMatcher
import requests
import re

MAX_PROMPTS_EACH_REQ = 70       # decrease this number if there is Error with AI model out of memory
MAX_WORDS_EACH_PROMPT = 50


# Backlogs: 
#   edit text with same format          Done
#   keep image and table                Done    
#   max_length                          Done
#   reference                           Done        (done it manually for both inline ref and ref cite) 
#   page break                          Done 
#   keep inline image                  Partially    (keep inline img but failed to recognize its postition)
#   suggestion text                      N/A        (will be removed from the docx file after grammar correct)
#   work with doc file                   N/A        (only accept docx file)

def similar(a, b):
    return SequenceMatcher(None, a, b).ratio()

# this function return an array of string as the result. If the array is empty -> failed to retrieve code
def promptAPI(prompts):
    API_URL = "https://polite-horribly-cub.ngrok-free.app/generate_code"
    # Note to ensure the response is more structure, the recommended prompt should follow this format:
    # [Correct english of this text: xxxxx. Here is the corrected version]

    max_prompts_lists = []    #store a list of max 80 prompts for each of element if num of prompts > 80 to avoid out of memory
    while len(prompts) > MAX_PROMPTS_EACH_REQ:
        max_prompts_lists.append(prompts[:MAX_PROMPTS_EACH_REQ])
        prompts = prompts[MAX_PROMPTS_EACH_REQ:]
    if prompts:
        max_prompts_lists.append(prompts[:])
        prompts.clear()

    result = []
    for max_prompts_list in max_prompts_lists:
        # Send the GET request with the prompts
        response = requests.get(API_URL, params=[("prompts", prompt[1]) for prompt in max_prompts_list])
        # Check the status code and response content
        if response.status_code == 200:
            generated_code_list = response.json()
            if len(generated_code_list)==1 and 'error' in generated_code_list:
                print('Failed with AI model out of memory')
                return []
            for _, code in enumerate(generated_code_list):
                # remove everything after '\n'
                code = code.split('\n',1)[0]
                # add to result
                result.append(code)
        else:
            print("Failed to retrieve response fron LLM API. Status code:", response.status_code)
            return []
    return result 

def AddRunWithFormat(pargraph, preserved_run, new_text):
    new_run = pargraph.add_run(new_text)
    new_run.font.name = preserved_run.font.name
    new_run.font.highlight_color = preserved_run.font.highlight_color
    new_run.font.size = preserved_run.font.size
    new_run.font.color.rgb = preserved_run.font.color.rgb
    new_run.font.color.theme_color = preserved_run.font.color.theme_color
    new_run.font.all_caps = preserved_run.font.all_caps
    new_run.font.complex_script = preserved_run.font.complex_script
    new_run.font.cs_bold = preserved_run.font.cs_bold
    new_run.font.cs_italic = preserved_run.font.cs_italic
    new_run.font.double_strike = preserved_run.font.double_strike
    new_run.font.emboss = preserved_run.font.emboss
    new_run.font.hidden = preserved_run.font.hidden
    new_run.font.highlight_color = preserved_run.font.highlight_color
    new_run.font.imprint = preserved_run.font.imprint
    new_run.font.math = preserved_run.font.math
    new_run.font.name = preserved_run.font.name
    new_run.font.no_proof = preserved_run.font.no_proof
    new_run.font.outline = preserved_run.font.outline
    new_run.font.rtl = preserved_run.font.rtl
    new_run.font.shadow = preserved_run.font.shadow
    new_run.font.size = preserved_run.font.size
    new_run.font.small_caps = preserved_run.font.small_caps
    new_run.font.snap_to_grid = preserved_run.font.snap_to_grid
    new_run.font.spec_vanish = preserved_run.font.spec_vanish
    new_run.font.strike = preserved_run.font.strike
    new_run.font.subscript = preserved_run.font.subscript
    new_run.font.superscript = preserved_run.font.superscript
    new_run.font.web_hidden = preserved_run.font.web_hidden
    new_run.font.bold = preserved_run.font.bold
    new_run.font.italic = preserved_run.font.italic
    new_run.font.underline = preserved_run.font.underline
    new_run.bold = preserved_run.bold
    new_run.italic = preserved_run.italic
    new_run.underline = preserved_run.underline
    # print(f'preserved_run: {preserved_run.text}*')
    # print(f'new_text: {new_text}*')

def AddPendingRuns(paragraph, changed_runs, pending_text):
    #print(f'num of pending runs: {len(changed_runs)}')
    
    # print(f'pending_text before: *{pending_text}*')
    for changed_run in changed_runs:
        # print(f'changed_run: {changed_run.text}')
        #print(f'changed run text: *{changed_run.text}*')
        # Finding a substring of words in pending text that has highest similarity with each run 
        if pending_text.startswith(" "):
            substring = " "
        else:
            substring = ""
        maxSimilar = 0
        for word in pending_text.split():
            substring += word + " "
            curSimilar = similar(changed_run.text, substring)
            if curSimilar > maxSimilar:
                maxSimilar = curSimilar

        if pending_text.startswith(" "):
            substring = " "
        else:
            substring = ""
        for word in pending_text.split():
            substring += word + " "
            curSimilar = similar(changed_run.text, substring)
            if maxSimilar == curSimilar:
                substring 
                break
        
        # if not pending_text.endswith(" "):
        substring = substring.removesuffix(" ")
        # print(f'substring: *{substring}*')
        # Add new run in paragraph using substring text and preservedly changed run format
        AddRunWithFormat(paragraph,changed_run,substring)
        # update pending text
        pending_text = pending_text.removeprefix(substring)
        #print(f'pending_text after: *{pending_text}*')
        # Add remaining pending text into paragraph and format as latest changed run if there is still pending text
    if pending_text:
        AddRunWithFormat(paragraph,changed_run, pending_text)

def hasImage(doc, r):
    """Checks if a run contains an image.
    
    :param run: The run object from docx.
    :return: image information if the run contains an image, None otherwise.
    """
    ns = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    }
    root = ET.fromstring(r._r.xml)
    if 'pic:pic' in r.element.xml:
        drawing = root.find('.//w:drawing', ns)
        blip = drawing.find('.//a:blip', ns)
        rId = blip.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        image = doc.part.related_parts[rId]
        image_bytes = image.blob
        filename = BytesIO(image_bytes)
        extent = drawing.find('.//wp:extent',ns)
        width = float(extent.attrib.get('cx'))
        height = float(extent.attrib.get('cy'))
        return (filename, width, height)
    return None

def replaceParagraph(doc, paragraph, paraIdx, correct_text):
    
    if paraIdx >= len(correct_text):
        print("paraIx > length of api_paragraphs")
        print(f"{paragraph.text}")
        return False
    # preserve all runs before grammar correct for format purpose
    preserved_runs = []
    for r in paragraph.runs:
        preserved_runs.append(r)
    
    paragraph.text = "" # remove all runs (also remove image)

    # add runs and format with text from api
    combined_run = ''
    changed_runs = [] # preserve any run changed by grammar correct
    whitespace_run = ""
    for r in preserved_runs:        
        img = hasImage(doc, r)
        if img:
            run = paragraph.add_run()
            run.add_picture(img[0],img[1], img[2])

        if not r.text:
            # check if there is page break
            if 'w:br' in r._element.xml and 'type="page"' in r._element.xml:
                if changed_runs: # handle all runs changed by LLM model before the page break
                    # Handle similar with the case "already met grammar correct in previous runs & no grammar correct in current run"
                    match = re.search(combined_run, correct_text[paraIdx])
                    if match != None:
                        match_string = match.group()
                        prefix = re.search(combined_run.removesuffix(".+"), match_string).group()
                        pending_text = match_string.removeprefix(prefix)
                        if pending_text:
                            AddPendingRuns(paragraph, changed_runs, pending_text)
                        changed_runs.clear()
                    else:
                        print(f'{combined_run}*')
                        print(f'correct text: {correct_text[paraIdx]}*')
                        print("Error changed_runs in replaceParagraph() func")
                        return False
                # add page break
                run = paragraph.add_run()
                run.add_break(WD_BREAK.PAGE)
            continue
            
        if r.text.isspace():    # handle some runs contain only space
            whitespace_run = " "
            continue
        if whitespace_run:  # handle some runs contain only space
            r.text = whitespace_run + r.text
            whitespace_run = ""
        combined_run += re.escape(r.text)   # combine_run is regex

        if len(r.text.split()) < 4:     #avoid the case '.+i' match with another place in the paragraph
            match = None    
        else:
            match = re.search(combined_run, correct_text[paraIdx])

        if match != None: # no grammar correct in current run
            if changed_runs:
                # Set text and format for previous runs that are changed by grammar correct
                match_string = match.group()
                    # Find pending text (this text was already grammar-corrected) before current run
                prefix = re.search(combined_run.removesuffix(re.escape(r.text)).removesuffix(".+"), match_string).group()
                pending_text = match_string.removesuffix(r.text).removeprefix(prefix)
                
                    # Add new runs into paragraph using text from pending_text and format from changed_runs
                if pending_text:
                    AddPendingRuns(paragraph, changed_runs, pending_text)
                    # Clear all pending runs cause they were handled
                changed_runs.clear()

            # Set text and format for current run (no grammar correct)
            AddRunWithFormat(paragraph,r,r.text)

        else:   # there is grammar correct in the current run
            # remove the current run text and add regex for next run match
            combined_run = combined_run.removesuffix(re.escape(r.text))
            if combined_run[len(combined_run) - 2:] != ".+":    # check last 2 characters
                combined_run += ".+"    # only allow '.+' 1 time although there may be multiple consecutive runs changed by grammar correct
            changed_runs.append(r)  # preserve changed run for formatting later

    # There are changed_runs (pending) at the end of paragraph
    if changed_runs:
        # Handle similar with the case "already met grammar correct in previous runs & no grammar correct in current run"
        match = re.search(combined_run, correct_text[paraIdx])
        if match != None:
            match_string = match.group()
            prefix = re.search(combined_run.removesuffix(".+"), match_string).group()
            pending_text = match_string.removeprefix(prefix)
            if pending_text:
                AddPendingRuns(paragraph, changed_runs, pending_text)
            changed_runs.clear()
        else:
            print(f'{combined_run}*')
            print(f'correct text: {correct_text[paraIdx]}*')
            print("Error changed_runs in replaceParagraph() func")
            return False

    return True

def editDocFile(document, api_paragraphs):
    # replace paragraphs in doc file with the result from API
    correct_idx = 0
    for p in document.paragraphs:
        if len(p.text.split()) < 5:
            if 'reference' in p.text.lower(): # reach reference section -> no need to edit more
                break
            continue
        if re.search("s[0-9]{7}", p.text) or ('.' not in p.text and '?' not in p.text):
            continue
        if replaceParagraph(document,p,correct_idx,api_paragraphs):
            correct_idx += 1
        else:
            print("Error in replaceParagraph() func")
            return False
        
    # replace tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if len(p.text.split()) < 5 or re.search("s[0-9]{7}", p.text) or ('.' not in p.text and '?' not in p.text):
                        continue
                    if replaceParagraph(document,p,correct_idx,api_paragraphs):
                        correct_idx += 1
                    else:
                        return False
    return True

def addPrompts(paragraph, para_idx, prompts, inline_references):
    if len(paragraph.text.split()) < 5:    # Not fixing any paragraph with less than 5 words cause it may be heading (may cause unexpected when fixing grammar for heading cause it not sentence) or space
        if 'reference' in paragraph.text.lower(): # reach reference section -> avoid prompting reference
            return False
        return True
    
    if re.search("s[0-9]{7}", paragraph.text): # ignore paragraph with student id 
        return True
    
    if '.' not in paragraph.text and '?' not in paragraph.text:   # ignore paragraph with no '.' (it may be heading, author name, title...)
        return True

    if '[' in paragraph.text: # may be inline references [] in paragraph 
        for sentence in re.split(r'(?<=[!?.])[ ]', paragraph.text):
            match = re.search("\[[0-9]+\]", sentence)
            if match:   # there is inline reference in format [] or ()
                inline_ref = match.group()
                prompts.append((para_idx, f"Correct english of this sentence: {sentence.replace(inline_ref, '')}\nHere is the corrected version:"))
                inline_references.append((len(prompts)-1,inline_ref))       
            else:   # no inline reference in sentence
                prompts.append((para_idx, f"Correct english of this sentence: {sentence}\nHere is the corrected version:"))
    elif '(' in paragraph.text:  # may be inline references () in paragraph
        uncompleted_sentences = ""  # in some cases reference contains '.' making a sentence not completed
        for sentence in re.split(r'(?<=[!?.])', paragraph.text):
            if re.search("et al\Z| p\Z|\(p\Z|,p\Z", sentence.removesuffix('.').rstrip()):
                uncompleted_sentences += sentence
                continue
            if uncompleted_sentences:
                completed_sentece = uncompleted_sentences + sentence
                uncompleted_sentences = ""
            else:
                completed_sentece = sentence
            match = re.search("\(.+\)$", completed_sentece.removesuffix('.').rstrip())
            if match:   # there is inline reference in format ()
                inline_ref = match.group()
                prompts.append((para_idx, f"Correct english of this sentence: {completed_sentece.replace(inline_ref, '')}\nHere is the corrected version:"))
                inline_references.append((len(prompts)-1,inline_ref))       
            else:   # () not at the end of sentence or () not exist in sentence -> no reference
                prompts.append((para_idx, f"Correct english of this sentence: {completed_sentece}\nHere is the corrected version:"))

                    
    else:    # no inline reference
        if len(paragraph.text.split()) > MAX_WORDS_EACH_PROMPT:  # check for word count
            combined_sentence = ""
            for sentence in re.split(r'(?<=[!?.])[ ]', paragraph.text):
                combined_sentence += sentence   #combine multiple sentences for prompting at once to reduce the number of prompts 
                if (len(combined_sentence.split()) > MAX_WORDS_EACH_PROMPT):   #max_lengh: 128 -> set max words to be 50 to preserve space for added words by AI
                    combined_sentence = combined_sentence.removesuffix(sentence)
                    prompts.append((para_idx, f"Correct english of this sentence: {combined_sentence}\nHere is the corrected version:"))
                    combined_sentence = sentence 
            if combined_sentence:
                prompts.append((para_idx, f"Correct english of this sentence: {combined_sentence}\nHere is the corrected version:"))
        else:
            prompts.append((para_idx, f"Correct english of this sentence: {paragraph.text}\nHere is the corrected version:"))
    return True

# only edit the docx file not copy to another one
# return True if successfully save the new doc
def fixDocGrammar(doc):
    # get the text in the doc file for prompting
    prompts = []    # prompt in form of (paragraphIndex, sentence)
    inline_references = []  #store inline reference to merge back into sentence later (avoiding AI makes change to inline references)
    # text in paragraphs (break into sentences if the paragraph word count > 50)
    for idx,p in enumerate(doc.paragraphs):
        # print(f'{idx}: *{p.text}*')       
        if not addPrompts(p,idx,prompts, inline_references): # reach reference section -> not add to prompt
            break

    #text in tables
    offset_idx = len(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    addPrompts(p,offset_idx,prompts,inline_references)
                    offset_idx += 1

    if not prompts: # doc file has no content
        doc.save('replace.docx')
        return True
    
    # call API to fix grammar
    api_response = promptAPI(prompts[:])
    if not api_response:    #something wrong in the response of AI
        return False

    if len(prompts) != len(api_response):
        print(f'Num of prompts: {len(prompts)}')
        print(f'Response length: {len(api_response)}')
        print("Something wrong in API response")
        return False
    
    # merge inline refenrences into sentences
    for prompt_idx,inline_ref in inline_references:
        api_response[prompt_idx] = ''.join([api_response[prompt_idx].rstrip().removesuffix('.'), ' ', inline_ref, '.'])

    api_paragraphs = []
    paragraph_text= ""
    paragraph_idx = prompts[0][0]
    for idx,response in enumerate(api_response):  
        if prompts[idx][0] > paragraph_idx:
            api_paragraphs.append(paragraph_text.rstrip())
            paragraph_text = ""
            paragraph_idx = prompts[idx][0]
        paragraph_text += response + " "
    api_paragraphs.append(paragraph_text.rstrip())

    if not editDocFile(doc, api_paragraphs):
        print("Error in editDocFile() func")
        return False

    return True

# # Function to modify the content of the DOCX file
# def modify_docx_content(file_path):
#     doc = Document(file_path)

#     # Start fixing grammar in docx file
#     if fixDocGrammar(doc):  # successfully
#         # Save the modified document
#         modified_file_path = file_path.replace('.docx', '_modified.docx')
#         doc.save(modified_file_path)
#         return modified_file_path
#     else:
#         print("Error during the doc file processing!!!")
#         return None

# source_file_path = 'Cultivating institutions for the national e-prescription.docx'
# print(modify_docx_content(source_file_path))



